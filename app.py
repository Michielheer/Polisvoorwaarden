import streamlit as st
import pdfplumber
import openai
from openai import OpenAI
from dotenv import load_dotenv
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import pandas as pd

# API key direct instellen
api_key = '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'
client = OpenAI(api_key=api_key)

st.set_page_config(page_title="PDF Vergelijker", layout="wide")
st.title("PDF Vergelijker")

# Upload bestanden
st.header("Upload je PDF's")
col1, col2 = st.columns(2)

with col1:
    pdf1 = st.file_uploader("Upload eerste polisvoorwaarden", type=['pdf'])
with col2:
    pdf2 = st.file_uploader("Upload tweede polisvoorwaarden", type=['pdf'])

def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text

def get_insurer_name(pdf_file):
    # Haal de bestandsnaam zonder .pdf extensie
    return os.path.splitext(pdf_file.name)[0]

def create_word_document(insurer1, insurer2, comparison_text, style=None):
    """
    Maak een professioneel Word-document van de vergelijking.
    - comparison_text: string (voor tekststijl) of lijst van dicts (voor tabelstijl)
    - style: 'tekst' of 'tabel' (optioneel, autodetect als None)
    """
    doc = Document()

    # Professionele titel
    title = doc.add_heading(f'Vergelijking polisvoorwaarden: {insurer1} vs {insurer2}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Datum
    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date.add_run(f'Datum: {datetime.datetime.now().strftime("%d-%m-%Y")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(120, 120, 120)

    # Ruimte
    doc.add_paragraph("")

    # Automatische stijlkeuze
    if style is None:
        if isinstance(comparison_text, list):
            style = 'tabel'
        else:
            style = 'tekst'

    if style == 'tabel' and isinstance(comparison_text, list):
        # Tabel met nette kopjes
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Onderdeel'
        hdr_cells[1].text = insurer1
        hdr_cells[2].text = insurer2
        hdr_cells[3].text = 'Status'
        hdr_cells[4].text = 'Details'
        for cell in hdr_cells:
            for p in cell.paragraphs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(10)
        # Rijen toevoegen
        for item in comparison_text:
            row = table.add_row().cells
            row[0].text = str(item.get('category', ''))
            row[1].text = str(item.get('insurer1', ''))
            row[2].text = str(item.get('insurer2', ''))
            row[3].text = str(item.get('status', ''))
            row[4].text = str(item.get('details', ''))
        doc.add_paragraph("")
    else:
        # Professionele tekstopmaak
        for blok in comparison_text.split('\n\n'):
            if blok.strip().startswith('- '):
                # Opsomming
                for regel in blok.strip().split('\n'):
                    doc.add_paragraph(regel.strip(), style='List Bullet')
            elif blok.strip().startswith('Artikel') or blok.strip().startswith('**'):
                p = doc.add_paragraph()
                run = p.add_run(blok.strip())
                run.bold = True
                run.font.size = Pt(12)
            else:
                p = doc.add_paragraph(blok.strip())
                p.paragraph_format.space_after = Pt(6)

    # Voetnoot
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ftr = footer.add_run('Deze vergelijking is automatisch gegenereerd en dient als indicatie. Raadpleeg altijd de originele polisvoorwaarden voor de exacte voorwaarden.')
    ftr.italic = True
    ftr.font.size = Pt(9)
    ftr.font.color.rgb = RGBColor(120, 120, 120)

    return doc

def parse_comparison_to_table(comparison_text, insurer1, insurer2):
    """
    Parse de vergelijkingstekst naar een tabel voor Streamlit
    """
    lines = comparison_text.split('\n')
    table_data = []
    current_category = ""
    current_insurer1_value = ""
    current_insurer2_value = ""
    current_difference = ""
    
    # Mogelijke categorieën die we kunnen herkennen
    categories = [
        "Artikel 1", "Artikel 2", "Artikel 3", "Artikel 4", "Artikel 5", "Artikel 6", "Artikel 7", "Artikel 8", "Artikel 9", "Artikel 10",
        "Eigen risico", "Maximumvergoedingen", "Dekking", "Uitsluitingen", "Niet verzekerd", "Niet gedekt", "Wat is niet verzekerd", "Niet meeverzekerd", "Beperkingen",
        "Verplichtingen", "Schade", "Dekking op locatie",
        "Gedekte schadeoorzaken", "Wat te doen bij schade", "Aanvullende dekkingen", "Verhuizing", "Verbouwing", "Verplichtingen bij schade"
    ]
    
    for line in lines:
        line = line.strip()
        
        # Check voor categorieën (verschillende formaten)
        if any(cat in line for cat in categories):
            # Nieuwe categorie gevonden - sla vorige op als die bestaat
            if current_category and (current_insurer1_value or current_insurer2_value):
                status = 'Verschil' if current_difference.strip() else 'Gelijk'
                table_data.append({
                    'Onderdeel': current_category,
                    insurer1: current_insurer1_value,
                    insurer2: current_insurer2_value,
                    'Verschil': current_difference,
                    'Status': status
                })
            
            # Reset voor nieuwe categorie
            for cat in categories:
                if cat in line:
                    current_category = cat
                    break
            current_insurer1_value = ""
            current_insurer2_value = ""
            current_difference = ""
            
        elif line.startswith('**') and line.endswith('**'):
            # Bold tekst als categorie
            current_category = line.replace('**', '').strip()
            current_insurer1_value = ""
            current_insurer2_value = ""
            current_difference = ""
            
        elif line.startswith('🔹'):
            # Verschil regel
            difference = line.replace('🔹', '').strip()
            
            if difference.startswith(f"{insurer1}:"):
                current_insurer1_value = difference.replace(f"{insurer1}:", "").strip()
            elif difference.startswith(f"{insurer2}:"):
                current_insurer2_value = difference.replace(f"{insurer2}:", "").strip()
            elif difference.startswith("Verschil:"):
                current_difference = difference.replace("Verschil:", "").strip()
    
    # Voeg laatste categorie toe
    if current_category and (current_insurer1_value or current_insurer2_value):
        status = 'Verschil' if current_difference.strip() else 'Gelijk'
        table_data.append({
            'Onderdeel': current_category,
            insurer1: current_insurer1_value,
            insurer2: current_insurer2_value,
            'Verschil': current_difference,
            'Status': status
        })
    
    # Als er nog geen data is, probeer een eenvoudigere aanpak
    if not table_data:
        # Zoek naar regels met verschillen
        for line in lines:
            line = line.strip()
            if line and ('verschil' in line.lower() or 'anders' in line.lower() or 'wel' in line.lower() or 'niet' in line.lower()):
                # Bepaal categorie op basis van context
                category = "Algemeen"
                for cat in categories:
                    if cat.lower() in comparison_text.lower():
                        category = cat
                        break
                
                table_data.append({
                    'Onderdeel': category,
                    insurer1: '✓' if insurer1.lower() in line.lower() else '',
                    insurer2: '✓' if insurer2.lower() in line.lower() else '',
                    'Verschil': line,
                    'Status': 'Verschil'
                })
    
    return pd.DataFrame(table_data)

def create_html_document(insurer1, insurer2, comparison_text, df):
    """
    Maak een professioneel HTML-document van de vergelijking met moderne styling.
    """
    # Parse de vergelijking naar tabeldata
    table_rows = ""
    if not df.empty:
        for _, row in df.iterrows():
            status_class = "status-different" if "verschil" in str(row['Status']).lower() else "status-same"
            status_text = "Verschil" if "verschil" in str(row['Status']).lower() else "Gelijk"
            table_rows += f"""
                <tr class="table-row">
                    <td class="px-6 py-3 whitespace-nowrap align-top">
                        <div class="text-sm font-medium text-gray-900">{row['Onderdeel']}</div>
                    </td>
                    <td class="px-6 py-3 align-top">
                        <div class="text-sm text-gray-700">{row[insurer1]}</div>
                    </td>
                    <td class="px-6 py-3 align-top">
                        <div class="text-sm text-gray-700">{row[insurer2]}</div>
                    </td>
                    <td class="px-6 py-3 whitespace-nowrap align-top">
                        <span class="status-badge {status_class}">{status_text}</span>
                    </td>
                    <td class="px-6 py-3 align-top">
                        <div class="text-sm text-gray-500 max-w-xs">{row['Verschil']}</div>
                    </td>
                </tr>
            """
    else:
        table_rows = """
            <tr class="table-row">
                <td colspan="5" class="px-6 py-3 text-center text-gray-500">
                    Geen gestructureerde verschillen gevonden
                </td>
            </tr>
        """

    html_content = f"""
    <!DOCTYPE html>
    <html lang="nl">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Vergelijking polisvoorwaarden: {insurer1} vs {insurer2}</title>
        <script src="https://cdn.tailwindcss.com"></script>
        <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap" rel="stylesheet">
        <style>
            body {{
                font-family: 'Inter', sans-serif;
                background: #f8fafb;
            }}
            .card-shadow {{
                box-shadow: 0 2px 8px 0 rgba(0,0,0,0.03);
            }}
            .status-badge {{
                padding: 0.15rem 0.7rem;
                border-radius: 9999px;
                font-size: 0.75rem;
                font-weight: 500;
                background: #f3f4f6;
                color: #6b7280;
                border: 1px solid #e5e7eb;
            }}
            .status-different {{
                background: #fef9c3;
                color: #a16207;
                border: 1px solid #fef3c7;
            }}
            .status-same {{
                background: #e0f2fe;
                color: #0369a1;
                border: 1px solid #bae6fd;
            }}
            .table-row {{
                transition: background 0.15s;
            }}
            .table-row:hover {{
                background: #f9fafb;
            }}
            .section-title {{
                font-size: 1.25rem;
                font-weight: 600;
                color: #22223b;
            }}
            .section-desc {{
                font-size: 1rem;
                color: #6b7280;
            }}
            .header-title {{
                font-size: 2rem;
                font-weight: 700;
                color: #22223b;
            }}
            .header-desc {{
                font-size: 1.1rem;
                color: #64748b;
            }}
            .comparison-text {{
                background: #f9fafb;
                border: 1.5px solid #e5e7eb;
                border-radius: 0.75rem;
                padding: 1rem 1.25rem;
                white-space: pre-wrap;
                line-height: 1.6;
                font-size: 0.875rem;
                color: #374151;
            }}
        </style>
    </head>
    <body>
        <!-- Header -->
        <header class="bg-white border-b border-gray-200 py-5 mb-8">
            <div class="max-w-5xl mx-auto px-4 flex flex-col gap-1">
                <span class="header-title">Vergelijking polisvoorwaarden</span>
                <span class="header-desc">{insurer1} vs {insurer2}</span>
            </div>
        </header>

        <main class="max-w-5xl mx-auto px-4">
            <!-- Comparison Results -->
            <section class="bg-white card-shadow rounded-xl overflow-hidden mb-8">
                <div class="px-6 py-4 border-b border-gray-100">
                    <div class="flex flex-col md:flex-row md:items-center md:justify-between gap-2">
                        <div>
                            <span class="section-title">Vergelijkingsresultaat</span>
                            <div class="section-desc mt-1">Overzicht van de verschillen tussen de polisvoorwaarden</div>
                        </div>
                        <div class="text-sm text-gray-500">
                            Datum: {datetime.datetime.now().strftime("%d-%m-%Y %H:%M")}
                        </div>
                    </div>
                </div>
                
                <!-- Summary Cards -->
                <div class="px-6 py-4 border-b border-gray-100">
                    <div class="grid grid-cols-2 md:grid-cols-4 gap-3">
                        <div class="bg-gray-50 border border-gray-200 rounded-lg p-3 flex items-center gap-2">
                            <div class="w-3 h-3 bg-yellow-200 rounded-full"></div>
                            <div>
                                <div class="text-base font-semibold">{len(df[df['Status'].str.contains('Verschil', case=False, na=False)]) if not df.empty else 0}</div>
                                <div class="text-xs text-gray-500">Verschillen</div>
                            </div>
                        </div>
                        <div class="bg-gray-50 border border-gray-200 rounded-lg p-3 flex items-center gap-2">
                            <div class="w-3 h-3 bg-blue-200 rounded-full"></div>
                            <div>
                                <div class="text-base font-semibold">{len(df[df['Status'].str.contains('Gelijk', case=False, na=False)]) if not df.empty else 0}</div>
                                <div class="text-xs text-gray-500">Overeenkomsten</div>
                            </div>
                        </div>
                        <div class="bg-gray-50 border border-gray-200 rounded-lg p-3 flex items-center gap-2">
                            <div class="w-3 h-3 bg-gray-300 rounded-full"></div>
                            <div>
                                <div class="text-base font-semibold">{len(df) if not df.empty else 0}</div>
                                <div class="text-xs text-gray-500">Totaal</div>
                            </div>
                        </div>
                        <div class="bg-gray-50 border border-gray-200 rounded-lg p-3 flex items-center gap-2">
                            <div class="w-3 h-3 bg-green-200 rounded-full"></div>
                            <div>
                                <div class="text-base font-semibold">100%</div>
                                <div class="text-xs text-gray-500">Compleet</div>
                            </div>
                        </div>
                    </div>
                </div>
                
                <!-- Table -->
                <div class="overflow-x-auto">
                    <table class="w-full">
                        <thead class="bg-gray-50">
                            <tr>
                                <th class="px-6 py-3 text-left text-xs font-medium text-gray-400 uppercase tracking-wider">Onderdeel</th>
                                <th class="px-6 py-3 text-left text-xs font-medium text-gray-400 uppercase tracking-wider">{insurer1}</th>
                                <th class="px-6 py-3 text-left text-xs font-medium text-gray-400 uppercase tracking-wider">{insurer2}</th>
                                <th class="px-6 py-3 text-left text-xs font-medium text-gray-400 uppercase tracking-wider">Status</th>
                                <th class="px-6 py-3 text-left text-xs font-medium text-gray-400 uppercase tracking-wider">Details</th>
                            </tr>
                        </thead>
                        <tbody class="bg-white divide-y divide-gray-100">
                            {table_rows}
                        </tbody>
                    </table>
                </div>
            </section>

            <!-- Detailed Comparison -->
            <section class="bg-white card-shadow rounded-xl p-6">
                <div class="mb-4">
                    <span class="section-title">Gedetailleerde vergelijking</span>
                    <div class="section-desc mt-1">Volledige analyse van alle verschillen en overeenkomsten</div>
                </div>
                <div class="comparison-text">{comparison_text}</div>
            </section>
        </main>

        <!-- Footer -->
        <footer class="bg-white border-t border-gray-200 py-6 mt-8">
            <div class="max-w-5xl mx-auto px-4 text-center">
                <p class="text-sm text-gray-500 italic">
                    Deze vergelijking is automatisch gegenereerd en dient als indicatie. 
                    Raadpleeg altijd de originele polisvoorwaarden voor de exacte voorwaarden.
                </p>
            </div>
        </footer>
    </body>
    </html>
    """
    return html_content

if pdf1 and pdf2:
    if st.button("Vergelijk documenten"):
        with st.spinner("Documenten worden vergeleken..."):
            # Extraheer tekst uit beide PDF's
            text1 = extract_text_from_pdf(pdf1)
            text2 = extract_text_from_pdf(pdf2)
            
            # Haal verzekeraarsnamen uit de bestandsnamen
            insurer1 = get_insurer_name(pdf1)
            insurer2 = get_insurer_name(pdf2)
            
            try:
                # Vraag GPT om de verschillen te analyseren
                response = client.chat.completions.create(
                    model="gpt-4o",
                    temperature=0.2,
                    messages=[
                        {
                            "role": "system",
                            "content": (
                                "Je bent een juridisch specialist in verzekeringsrecht. Je vergelijkt polisvoorwaarden voor verzekeringen "
                                "en geeft **concrete, specifieke verschillen** met exacte bedragen, voorwaarden en uitsluitingen. "
                                "Neem bedragen, voorwaarden en uitsluitingen letterlijk en identiek over uit de polis. "
                                "Geef alleen een verschil als het bedrag of de tekst écht anders is. "
                                "Let extra goed op uitsluitingen, ook als deze onder kopjes als 'Uitsluitingen', 'Niet verzekerd', 'Niet gedekt', 'Wat is niet verzekerd', 'Niet meeverzekerd', 'Beperkingen', of soortgelijke termen staan. "
                                "Gebruik altijd het format met **categorie** en 🔹 voor verschillen."
                            )
                        },
                        {
                            "role": "user",
                            "content": (
                                f"Vergelijk de polisvoorwaarden van {insurer1} en {insurer2}. "
                                "Geef voor elk onderdeel de **exacte bepalingen** van beide verzekeraars, letterlijk overgenomen uit de polis. "
                                "Geef alleen een verschil als het bedrag, de voorwaarde of de uitsluiting écht anders is. "
                                "Laat bedragen, voorwaarden en uitsluitingen exact zien zoals ze in de polis staan.\n\n"
                                "Let extra goed op uitsluitingen, ook als deze onder kopjes als 'Uitsluitingen', 'Niet verzekerd', 'Niet gedekt', 'Wat is niet verzekerd', 'Niet meeverzekerd', 'Beperkingen', of soortgelijke termen staan.\n\n"
                                "**Gebruik dit exacte format per onderdeel:**\n"
                                "**Artikel 1: Gedekte schadeoorzaken**\n"
                                f"🔹 {insurer1}: [letterlijke tekst uit document]\n"
                                f"🔹 {insurer2}: [letterlijke tekst uit document]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                "**Eigen risico**\n"
                                f"🔹 {insurer1}: [letterlijk bedrag/percentage]\n"
                                f"🔹 {insurer2}: [letterlijk bedrag/percentage]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                "**Maximumvergoedingen**\n"
                                f"🔹 {insurer1}: [letterlijk bedrag]\n"
                                f"🔹 {insurer2}: [letterlijk bedrag]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                "**Uitsluitingen / Niet verzekerd / Niet gedekt / Wat is niet verzekerd / Niet meeverzekerd / Beperkingen**\n"
                                f"🔹 {insurer1}: [letterlijke uitsluitingen]\n"
                                f"🔹 {insurer2}: [letterlijke uitsluitingen]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                "**Dekking op locatie**\n"
                                f"🔹 {insurer1}: [letterlijke dekking]\n"
                                f"🔹 {insurer2}: [letterlijke dekking]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                "**Verplichtingen bij schade**\n"
                                f"🔹 {insurer1}: [letterlijke verplichtingen]\n"
                                f"🔹 {insurer2}: [letterlijke verplichtingen]\n"
                                f"🔹 Verschil: [alleen invullen als er écht een verschil is]\n\n"
                                f"{insurer1}:\n{text1}\n\n"
                                f"{insurer2}:\n{text2}"
                            )
                        }
                    ]
                )
                
                comparison_text = response.choices[0].message.content
                
                # Toon resultaten
                st.header(f"Verschillen tussen {insurer1} en {insurer2}")
                
                # Toon als tekst
                st.subheader("📄 Gedetailleerde vergelijking")
                st.write(comparison_text)
                
                # Toon als tabel
                st.subheader("📊 Overzichtstabel")
                try:
                    df = parse_comparison_to_table(comparison_text, insurer1, insurer2)
                    if not df.empty:
                        st.dataframe(df, use_container_width=True)
                    else:
                        st.info("Geen gestructureerde verschillen gevonden voor tabelweergave.")
                except Exception as e:
                    st.warning(f"Kon tabel niet genereren: {str(e)}")
                
                # Maak Word document
                doc = create_word_document(insurer1, insurer2, comparison_text)
                
                # Sla document op
                filename = f"Vergelijking_{insurer1}_vs_{insurer2}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                doc.save(filename)
                
                # Maak HTML document
                html_content = create_html_document(insurer1, insurer2, comparison_text, df)
                html_filename = f"Vergelijking_{insurer1}_vs_{insurer2}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
                
                with open(html_filename, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                # Download knoppen
                col1, col2 = st.columns(2)
                with col1:
                    with open(filename, 'rb') as f:
                        st.download_button(
                            label="📥 Download Word document",
                            data=f,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                
                with col2:
                    with open(html_filename, 'r', encoding='utf-8') as f:
                        st.download_button(
                            label="🌐 Download HTML document",
                            data=f.read(),
                            file_name=html_filename,
                            mime="text/html"
                        )
                
            except Exception as e:
                st.error(f"Er is een fout opgetreden: {str(e)}")
                st.info("Controleer of de API key correct is ingesteld.")
